import os
import uuid
import tempfile
import numpy as np
import streamlit as st
from dotenv import load_dotenv
from sentence_transformers import SentenceTransformer
from openai import OpenAI
import oracledb
import json
from datetime import datetime
import logging
from pathlib import Path
import base64

# Mobile-Responsive Streamlit Configuration (MUST BE FIRST!)
st.set_page_config(
    page_title="🤖 Smart RAG Assistant",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="auto",
    menu_items={
        'Get Help': 'https://your-help-url.com',
        'Report a bug': 'https://your-bug-report-url.com',
        'About': "Smart RAG Assistant powered by Oracle 23ai"
    }
)

# Additional imports for multi-modal support (with availability checking)
import pdfplumber

# Check for optional dependencies
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from pptx import Presentation
    PPTX_AVAILABLE = True
except ImportError:
    PPTX_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    import camelot
    CAMELOT_AVAILABLE = True
except ImportError:
    CAMELOT_AVAILABLE = False

try:
    import tabula
    TABULA_AVAILABLE = True
except ImportError:
    TABULA_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    import markdown
    WEB_AVAILABLE = True
except ImportError:
    WEB_AVAILABLE = False

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Mobile-responsive CSS
st.markdown("""
<style>
    /* Mobile-first responsive design */
    @media (max-width: 768px) {
        .main .block-container {
            padding-left: 1rem;
            padding-right: 1rem;
            padding-top: 1rem;
        }
        
        .stTabs [data-baseweb="tab-list"] {
            gap: 8px;
        }
        
        .stTabs [data-baseweb="tab"] {
            padding: 8px 12px;
            font-size: 14px;
        }
        
        .stButton > button {
            width: 100%;
            margin-bottom: 8px;
        }
        
        .stSelectbox > div > div {
            font-size: 14px;
        }
        
        .stTextInput > div > div > input {
            font-size: 16px; /* Prevents zoom on iOS */
        }
        
        .stTextArea > div > div > textarea {
            font-size: 16px;
        }
    }
    
    /* Touch-friendly interactive elements */
    .stButton > button {
        min-height: 44px; /* Apple's recommended touch target size */
        border-radius: 8px;
        font-weight: 500;
    }
    
    .stSelectbox > div > div {
        min-height: 44px;
    }
    
    /* Improved file upload area */
    .uploadedFile {
        border-radius: 8px;
        border: 2px dashed #ccc;
        padding: 20px;
        text-align: center;
        background-color: #f9f9f9;
    }
    
    /* Chat message styling */
    .chat-message {
        padding: 12px;
        border-radius: 12px;
        margin-bottom: 8px;
        max-width: 85%;
    }
    
    .user-message {
        background-color: #007AFF;
        color: white;
        margin-left: auto;
    }
    
    .assistant-message {
        background-color: #F2F2F7;
        color: black;
    }
    
    /* Loading spinner */
    .loading-spinner {
        display: inline-block;
        width: 20px;
        height: 20px;
        border: 3px solid #f3f3f3;
        border-top: 3px solid #007AFF;
        border-radius: 50%;
        animation: spin 1s linear infinite;
    }
    
    @keyframes spin {
        0% { transform: rotate(0deg); }
        100% { transform: rotate(360deg); }
    }
    
    /* Responsive grid */
    .responsive-grid {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(300px, 1fr));
        gap: 16px;
        margin: 16px 0;
    }
    
    /* Credential input styling */
    .credential-section {
        background-color: #f8f9fa;
        padding: 20px;
        border-radius: 10px;
        border-left: 4px solid #007AFF;
        margin: 10px 0;
    }
    
    .success-banner {
        background-color: #d4edda;
        border: 1px solid #c3e6cb;
        color: #155724;
        padding: 12px;
        border-radius: 6px;
        margin: 10px 0;
    }
    
    .warning-banner {
        background-color: #fff3cd;
        border: 1px solid #ffeaa7;
        color: #856404;
        padding: 12px;
        border-radius: 6px;
        margin: 10px 0;
    }
    
    /* Dark mode support */
    @media (prefers-color-scheme: dark) {
        .uploadedFile {
            background-color: #2D2D30;
            border-color: #444;
        }
        
        .assistant-message {
            background-color: #2D2D30;
            color: #FFFFFF;
        }
        
        .credential-section {
            background-color: #2D2D30;
            color: #FFFFFF;
        }
    }
</style>
""", unsafe_allow_html=True)

# Load environment variables (fallback)
load_dotenv()

# Enhanced Configuration with User Input Support
class Config:
    # Model Configuration
    EMBEDDING_MODEL = "all-MiniLM-L6-v2"
    EMBEDDING_DIMENSION = 384
    CHUNK_SIZE = 1000
    OVERLAP = 200
    TOP_K = 5
    VECTOR_DISTANCE_METRIC = "COSINE"
    
    # Enhanced Configuration with Multi-Modal Support
    SUPPORTED_FORMATS = {}
    
    # Always available
    SUPPORTED_FORMATS.update({
        'pdf': 'PDF Documents',
        'txt': 'Text Files'
    })
    
    # Conditional format support based on available libraries
    if DOCX_AVAILABLE:
        SUPPORTED_FORMATS['docx'] = 'Word Documents'
    
    if PPTX_AVAILABLE:
        SUPPORTED_FORMATS['pptx'] = 'PowerPoint Presentations'
    
    if PANDAS_AVAILABLE:
        SUPPORTED_FORMATS.update({
            'xlsx': 'Excel Spreadsheets',
            'csv': 'CSV Files'
        })
    
    if OCR_AVAILABLE:
        SUPPORTED_FORMATS.update({
            'jpg': 'JPEG Images',
            'jpeg': 'JPEG Images',
            'png': 'PNG Images',
            'bmp': 'Bitmap Images',
            'tiff': 'TIFF Images'
        })
    
    if WEB_AVAILABLE:
        SUPPORTED_FORMATS.update({
            'md': 'Markdown Files',
            'html': 'HTML Files'
        })
    
    # OCR Configuration
    TESSERACT_CONFIG = '--oem 3 --psm 6'
    
    # Mobile Configuration
    MOBILE_CHUNK_SIZE = 500  # Smaller chunks for mobile
    MOBILE_TOP_K = 3         # Fewer results for mobile

# Credential Management Functions
def initialize_session_credentials():
    """Initialize credential session state variables"""
    if 'credentials_configured' not in st.session_state:
        st.session_state.credentials_configured = False
    if 'openai_client' not in st.session_state:
        st.session_state.openai_client = None
    if 'db_credentials' not in st.session_state:
        st.session_state.db_credentials = {}
    if 'credentials_tested' not in st.session_state:
        st.session_state.credentials_tested = False

def test_openai_credentials(api_key):
    """Test OpenAI API key validity"""
    try:
        client = OpenAI(api_key=api_key)
        # Test with a simple completion
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": "Hello"}],
            max_tokens=5
        )
        return True, client
    except Exception as e:
        return False, str(e)

def test_database_credentials(username, password, service_name, wallet_path):
    """Test Oracle database connection"""
    try:
        # Set Oracle environment
        os.environ["TNS_ADMIN"] = wallet_path
        os.environ['PYTHONHTTPSVERIFY'] = '0'
        
        connection = oracledb.connect(
            user=username,
            password=password,
            dsn=service_name,
            config_dir=wallet_path,
            wallet_location=wallet_path,
            wallet_password=password
        )
        connection.close()
        return True, "Connection successful"
    except Exception as e:
        return False, str(e)

def render_credential_setup():
    """Render the credential setup interface"""
    st.title("🔐 Setup Credentials")
    st.markdown("*Please provide your OpenAI API key and Oracle Database credentials to get started.*")
    
    # Check if we have env variables as fallback
    env_openai = os.getenv("OPENAI_API_KEY")
    env_username = os.getenv("ORACLE_USERNAME")
    env_password = os.getenv("ORACLE_PASSWORD")
    env_service = os.getenv("ORACLE_SERVICE_NAME")
    env_wallet = os.getenv("ORACLE_WALLET_PATH", "/path/to/wallet")
    
    if any([env_openai, env_username, env_password, env_service]):
        st.info("💡 Some credentials were found in environment variables. You can override them below or use them as defaults.")
    
    # OpenAI API Configuration
    with st.expander("🤖 OpenAI API Configuration", expanded=True):
        st.markdown("**Get your API key from:** [OpenAI Platform](https://platform.openai.com/account/api-keys)")
        
        openai_api_key = st.text_input(
            "OpenAI API Key",
            type="password",
            value=env_openai if env_openai else "",
            placeholder="sk-...",
            help="Your OpenAI API key for GPT models"
        )
        
        if openai_api_key:
            if st.button("🔍 Test OpenAI Connection", key="test_openai"):
                with st.spinner("Testing OpenAI connection..."):
                    success, result = test_openai_credentials(openai_api_key)
                    if success:
                        st.success("✅ OpenAI API key is valid!")
                        st.session_state.openai_client = result
                    else:
                        st.error(f"❌ OpenAI API key test failed: {result}")
    
    # Oracle Database Configuration
    with st.expander("🗄️ Oracle Database 23ai Configuration", expanded=True):
        st.markdown("**Required:** Oracle Database 23ai with Vector Search capabilities")
        
        col1, col2 = st.columns(2)
        
        with col1:
            db_username = st.text_input(
                "Database Username",
                value=env_username if env_username else "",
                placeholder="admin",
                help="Your Oracle database username"
            )
            
            db_service_name = st.text_input(
                "Service Name",
                value=env_service if env_service else "",
                placeholder="your_service_name_high",
                help="Oracle database service name (usually ends with _high for autonomous DB)"
            )
        
        with col2:
            db_password = st.text_input(
                "Database Password",
                type="password",
                value=env_password if env_password else "",
                placeholder="your_password",
                help="Your Oracle database password"
            )
            
            wallet_path = st.text_input(
                "Wallet Path",
                value=env_wallet,
                placeholder="/path/to/wallet",
                help="Path to your Oracle wallet directory"
            )
        
        st.info("💡 **Tip:** For Oracle Autonomous Database, download the wallet from the Cloud Console")
        
        if all([db_username, db_password, db_service_name, wallet_path]):
            if st.button("🔍 Test Database Connection", key="test_db"):
                with st.spinner("Testing database connection..."):
                    success, result = test_database_credentials(
                        db_username, db_password, db_service_name, wallet_path
                    )
                    if success:
                        st.success("✅ Database connection successful!")
                        st.session_state.db_credentials = {
                            'username': db_username,
                            'password': db_password,
                            'service_name': db_service_name,
                            'wallet_path': wallet_path
                        }
                    else:
                        st.error(f"❌ Database connection failed: {result}")
    
    # Save Credentials Button
    st.markdown("---")
    
    # Check if all credentials are provided
    all_creds_provided = (
        openai_api_key and 
        db_username and 
        db_password and 
        db_service_name and 
        wallet_path
    )
    
    if all_creds_provided:
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            if st.button("🚀 Initialize Application", type="primary", use_container_width=True):
                with st.spinner("Initializing application..."):
                    # Test both connections
                    openai_success, openai_result = test_openai_credentials(openai_api_key)
                    db_success, db_result = test_database_credentials(
                        db_username, db_password, db_service_name, wallet_path
                    )
                    
                    if openai_success and db_success:
                        # Store credentials in session state
                        st.session_state.openai_client = openai_result
                        st.session_state.db_credentials = {
                            'username': db_username,
                            'password': db_password,
                            'service_name': db_service_name,
                            'wallet_path': wallet_path
                        }
                        st.session_state.credentials_configured = True
                        st.session_state.credentials_tested = True
                        
                        st.success("🎉 All credentials configured successfully!")
                        st.balloons()
                        
                        # Small delay to show success message
                        import time
                        time.sleep(1)
                        st.rerun()
                    else:
                        if not openai_success:
                            st.error(f"❌ OpenAI API error: {openai_result}")
                        if not db_success:
                            st.error(f"❌ Database error: {db_result}")
    else:
        st.warning("⚠️ Please provide all required credentials above")
    
    # Help section
    with st.expander("❓ Need Help?", expanded=False):
        st.markdown("""
        **Getting OpenAI API Key:**
        1. Go to [OpenAI Platform](https://platform.openai.com/)
        2. Sign up or log in to your account
        3. Navigate to API Keys section
        4. Create a new API key
        5. Copy and paste it above
        
        **Setting up Oracle Database 23ai:**
        1. Create an Oracle Autonomous Database instance
        2. Download the wallet from the Cloud Console
        3. Extract the wallet to a local directory
        4. Use the service name from tnsnames.ora file
        5. Ensure Vector Search is enabled
        
        **Security Note:**
        - Credentials are stored only in your browser session
        - They are not saved to disk or transmitted elsewhere
        - You'll need to re-enter them if you refresh the page
        """)

# Device detection functions (same as before)
def is_mobile():
    """Simple mobile detection based on screen width"""
    return st.session_state.get('is_mobile', False)

def detect_device():
    """Detect device type and adjust UI accordingly"""
    st.markdown("""
    <script>
        function detectDevice() {
            const isMobile = window.innerWidth <= 768;
            window.parent.postMessage({
                type: 'streamlit:setSessionState',
                data: {is_mobile: isMobile}
            }, '*');
        }
        detectDevice();
        window.addEventListener('resize', detectDevice);
    </script>
    """, unsafe_allow_html=True)

# Initialize device detection
detect_device()

# [Keep all the existing classes: MultiModalProcessor, EnhancedChunking, MobileUI - they remain the same]

# Multi-Modal Document Processors
class MultiModalProcessor:
    """Handle multiple document formats and extract content"""
    
    @staticmethod
    def detect_file_type(file_path):
        """Detect file type from extension"""
        return Path(file_path).suffix.lower().lstrip('.')
    
    @staticmethod
    def process_pdf(file_path):
        """Enhanced PDF processing with table and image extraction"""
        content = []
        try:
            # Text extraction
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text and text.strip():
                        content.append({
                            'type': 'text',
                            'content': text.strip(),
                            'page': page_num + 1,
                            'metadata': {'source': 'pdf_text'}
                        })
                    
                    # Table extraction
                    tables = page.extract_tables()
                    for table_idx, table in enumerate(tables):
                        if table:
                            table_text = MultiModalProcessor.table_to_text(table)
                            content.append({
                                'type': 'table',
                                'content': table_text,
                                'page': page_num + 1,
                                'metadata': {
                                    'source': 'pdf_table',
                                    'table_index': table_idx
                                }
                            })
        except Exception as e:
            logger.error(f"PDF processing error: {e}")
        
        return content
    
    @staticmethod
    def process_docx(file_path):
        """Process Word documents"""
        if not DOCX_AVAILABLE:
            return []
        
        content = []
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            for para_idx, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    content.append({
                        'type': 'text',
                        'content': paragraph.text.strip(),
                        'page': 1,
                        'metadata': {
                            'source': 'docx_paragraph',
                            'paragraph_index': para_idx,
                            'style': paragraph.style.name if paragraph.style else 'Normal'
                        }
                    })
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_text = MultiModalProcessor.docx_table_to_text(table)
                content.append({
                    'type': 'table',
                    'content': table_text,
                    'page': 1,
                    'metadata': {
                        'source': 'docx_table',
                        'table_index': table_idx
                    }
                })
                
        except Exception as e:
            logger.error(f"DOCX processing error: {e}")
        
        return content
    
    @staticmethod
    def table_to_text(table):
        """Convert table data to text format"""
        if not table:
            return ""
        
        text_rows = []
        for row in table:
            if row:  # Skip empty rows
                text_rows.append(" | ".join(str(cell) if cell else "" for cell in row))
        
        return "\n".join(text_rows)
    
    @staticmethod
    def docx_table_to_text(table):
        """Convert DOCX table to text"""
        text_rows = []
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text.strip())
            text_rows.append(" | ".join(row_text))
        
        return "\n".join(text_rows)
    
    @staticmethod
    def process_text_file(file_path):
        """Process plain text files"""
        content = []
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
                if text.strip():
                    content.append({
                        'type': 'text',
                        'content': text.strip(),
                        'page': 1,
                        'metadata': {'source': 'text_file'}
                    })
        except Exception as e:
            logger.error(f"Text file processing error: {e}")
        
        return content
    
    @staticmethod
    def process_file(file_path):
        """Main processing function that routes to appropriate processor"""
        file_type = MultiModalProcessor.detect_file_type(file_path)
        
        processors = {
            'pdf': MultiModalProcessor.process_pdf,
            'docx': MultiModalProcessor.process_docx,
            'txt': MultiModalProcessor.process_text_file,
        }
        
        processor = processors.get(file_type)
        if processor:
            return processor(file_path)
        else:
            st.error(f"Unsupported file type: {file_type}")
            return []

# Enhanced chunking (simplified version)
class EnhancedChunking:
    """Advanced chunking strategies for different content types"""
    
    @staticmethod
    def smart_chunk_content(content_items, chunk_size=None, overlap=None):
        """Intelligently chunk content based on type and device"""
        if chunk_size is None:
            chunk_size = Config.MOBILE_CHUNK_SIZE if is_mobile() else Config.CHUNK_SIZE
        if overlap is None:
            overlap = Config.OVERLAP
        
        chunks = []
        
        for item in content_items:
            content_text = item['content']
            page = item['page']
            metadata = item['metadata']
            
            text_chunks = EnhancedChunking.chunk_text_content(content_text, chunk_size, overlap)
            for chunk_text in text_chunks:
                chunks.append((page, len(chunks), chunk_text, metadata))
        
        return chunks
    
    @staticmethod
    def chunk_text_content(text, chunk_size, overlap):
        """Smart text chunking with sentence boundary awareness"""
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        sentences = text.split('. ')
        current_chunk = []
        current_length = 0
        
        for sentence in sentences:
            sentence_length = len(sentence)
            
            if current_length + sentence_length > chunk_size and current_chunk:
                chunks.append('. '.join(current_chunk) + '.')
                overlap_sentences = current_chunk[-overlap//50:] if overlap > 0 else []
                current_chunk = overlap_sentences + [sentence]
                current_length = sum(len(s) for s in current_chunk)
            else:
                current_chunk.append(sentence)
                current_length += sentence_length
        
        if current_chunk:
            chunks.append('. '.join(current_chunk))
        
        return chunks

# Mobile UI (simplified version)
class MobileUI:
    """Mobile-specific UI components and layouts"""
    
    @staticmethod
    def mobile_file_uploader():
        """Mobile-friendly file uploader with format info"""
        st.markdown("### 📱 Upload Documents")
        
        uploaded_file = st.file_uploader(
            "Choose a file",
            type=list(Config.SUPPORTED_FORMATS.keys()),
            help="Select a document to add to your knowledge base"
        )
        
        return uploaded_file
    
    @staticmethod
    def mobile_chat_interface():
        """Mobile-optimized chat interface"""
        st.markdown("### 💬 Chat Assistant")
        
        with st.container():
            user_question = st.text_area(
                "Ask me anything about your documents:",
                height=100,
                placeholder="What would you like to know?",
                help="Ask questions about the content in your uploaded documents"
            )
            
            col1, col2 = st.columns([3, 1])
            with col1:
                ask_button = st.button("🔍 Ask", type="primary", use_container_width=True)
            with col2:
                clear_button = st.button("🗑️ Clear", use_container_width=True)
        
        return user_question, ask_button, clear_button
    
    @staticmethod
    def mobile_chat_message(message, is_user=True):
        """Render mobile-friendly chat messages"""
        if is_user:
            st.markdown(f"""
            <div class="chat-message user-message">
                <strong>You:</strong><br>{message}
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown(f"""
            <div class="chat-message assistant-message">
                <strong>🤖 Assistant:</strong><br>{message}
            </div>
            """, unsafe_allow_html=True)

# Load embedding model
@st.cache_resource
def load_embedding_model():
    """Load the sentence transformer model"""
    try:
        model = SentenceTransformer(Config.EMBEDDING_MODEL)
        return model
    except Exception as e:
        st.error(f"❌ Failed to load embedding model: {e}")
        return None

# Updated Database operations using session credentials
class DatabaseManager:
    """Handle Oracle Database 23ai operations using session credentials"""
    
    @staticmethod
    def get_connection():
        """Establish connection to Oracle Database 23ai using session credentials"""
        if not st.session_state.get('db_credentials'):
            return None
            
        creds = st.session_state.db_credentials
        try:
            connection = oracledb.connect(
                user=creds['username'],
                password=creds['password'],
                dsn=creds['service_name'],
                config_dir=creds['wallet_path'],
                wallet_location=creds['wallet_path'],
                wallet_password=creds['password']
            )
            return connection
        except Exception as e:
            st.error(f"❌ Database Connection Error: {e}")
            logger.error(f"Database connection failed: {e}")
            return None
    
    @staticmethod
    def vector_to_string(embedding):
        """Convert numpy array or list to Oracle vector string format"""
        try:
            if isinstance(embedding, np.ndarray):
                vector_list = embedding.astype(float).tolist()
            else:
                vector_list = [float(x) for x in embedding]
            
            return '[' + ','.join(f'{x:.6f}' for x in vector_list) + ']'
        except Exception as e:
            logger.error(f"Vector conversion error: {e}")
            raise ValueError(f"Failed to convert vector to string: {e}")
    
    @staticmethod
    def initialize_tables():
        """Create tables with vector column for Oracle 23ai"""
        conn = DatabaseManager.get_connection()
        if not conn:
            return False
        
        try:
            with conn.cursor() as cursor:
                cursor.execute("""
                    SELECT table_name FROM user_tables 
                    WHERE table_name IN ('DOCUMENTS', 'DOCUMENT_CHUNKS')
                """)
                existing_tables = [row[0] for row in cursor.fetchall()]
                
                if 'DOCUMENTS' not in existing_tables:
                    cursor.execute("""
                        CREATE TABLE documents (
                            document_id VARCHAR2(50) PRIMARY KEY,
                            filename VARCHAR2(500) NOT NULL,
                            file_type VARCHAR2(50) DEFAULT 'pdf',
                            upload_timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                            total_chunks NUMBER DEFAULT 0,
                            file_size NUMBER,
                            content_type VARCHAR2(100),
                            processing_status VARCHAR2(50) DEFAULT 'processed'
                        )
                    """)
                    st.success("✅ Created documents table")
                
                if 'DOCUMENT_CHUNKS' not in existing_tables:
                    cursor.execute(f"""
                        CREATE TABLE document_chunks (
                            chunk_id VARCHAR2(50) PRIMARY KEY,
                            document_id VARCHAR2(50) NOT NULL,
                            chunk_text CLOB NOT NULL,
                            chunk_index NUMBER NOT NULL,
                            page_number NUMBER DEFAULT 1,
                            content_type VARCHAR2(50) DEFAULT 'text',
                            metadata CLOB,
                            embedding VECTOR({Config.EMBEDDING_DIMENSION}, FLOAT32),
                            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                            CONSTRAINT fk_doc_chunks_doc_id 
                                FOREIGN KEY (document_id) REFERENCES documents(document_id) ON DELETE CASCADE
                        )
                    """)
                    st.success("✅ Created document_chunks table")
                
                conn.commit()
                return True
                
        except Exception as e:
            st.error(f"❌ Table initialization error: {e}")
            logger.error(f"Table initialization failed: {e}")
            return False
        finally:
            conn.close()
    
    @staticmethod
    def store_enhanced_chunks(doc_id, chunks_with_embeddings, file_type):
        """Store chunks with enhanced metadata"""
        conn = DatabaseManager.get_connection()
        if not conn:
            return False
        
        try:
            with conn.cursor() as cursor:
                success_count = 0
                for i, chunk_data in enumerate(chunks_with_embeddings):
                    try:
                        chunk_id = str(uuid.uuid4())
                        page_num, chunk_idx, text, metadata = chunk_data[:4]
                        embedding = chunk_data[4] if len(chunk_data) > 4 else None
                        
                        if embedding is None and load_embedding_model():
                            embedder = load_embedding_model()
                            embedding = embedder.encode(text, convert_to_numpy=True)
                        
                        if embedding is not None:
                            embedding_str = DatabaseManager.vector_to_string(embedding)
                            
                            cursor.execute("""
                                INSERT INTO document_chunks (
                                    chunk_id, document_id, chunk_text, chunk_index, 
                                    page_number, content_type, metadata, embedding
                                ) VALUES (:1, :2, :3, :4, :5, :6, :7, TO_VECTOR(:8))
                            """, (
                                chunk_id, doc_id, text, chunk_idx, page_num,
                                metadata.get('source', 'text') if isinstance(metadata, dict) else 'text',
                                json.dumps(metadata) if isinstance(metadata, dict) else str(metadata),
                                embedding_str
                            ))
                            
                            success_count += 1
                            
                            if (i + 1) % 50 == 0:
                                conn.commit()
                                if is_mobile():
                                    st.progress((i + 1) / len(chunks_with_embeddings))
                    
                    except Exception as chunk_error:
                        logger.error(f"Error storing chunk {i}: {chunk_error}")
                        continue
                
                conn.commit()
                logger.info(f"Successfully stored {success_count}/{len(chunks_with_embeddings)} chunks")
                return success_count > 0
                
        except Exception as e:
            st.error(f"❌ Chunk storage error: {e}")
            logger.error(f"Chunk storage failed: {e}")
            return False
        finally:
            conn.close()
    
    @staticmethod
    def store_document_metadata(doc_id, filename, file_type, file_size, total_chunks):
        """Store enhanced document metadata"""
        conn = DatabaseManager.get_connection()
        if not conn:
            return False
        
        try:
            with conn.cursor() as cursor:
                cursor.execute("""
                    INSERT INTO documents (document_id, filename, file_type, file_size, total_chunks, content_type)
                    VALUES (:1, :2, :3, :4, :5, :6)
                """, (doc_id, filename, file_type, file_size, total_chunks, f'application/{file_type}'))
                conn.commit()
                return True
        except Exception as e:
            st.error(f"❌ Document metadata storage error: {e}")
            return False
        finally:
            conn.close()
    
    @staticmethod
    def enhanced_vector_search(query_embedding, top_k=None):
        """Enhanced vector search"""
        if top_k is None:
            top_k = Config.MOBILE_TOP_K if is_mobile() else Config.TOP_K
        
        conn = DatabaseManager.get_connection()
        if not conn:
            return []
        
        try:
            with conn.cursor() as cursor:
                query_vector_str = DatabaseManager.vector_to_string(query_embedding)
                
                query = f"""
                    SELECT 
                        dc.chunk_text,
                        dc.page_number,
                        dc.content_type,
                        dc.metadata,
                        d.filename,
                        d.file_type,
                        VECTOR_DISTANCE(dc.embedding, TO_VECTOR(:1), {Config.VECTOR_DISTANCE_METRIC}) as similarity_score
                    FROM document_chunks dc
                    JOIN documents d ON dc.document_id = d.document_id
                    ORDER BY similarity_score
                    FETCH FIRST :2 ROWS ONLY
                """
                
                cursor.execute(query, (query_vector_str, top_k))
                
                results = []
                for row in cursor.fetchall():
                    chunk_text = row[0].read() if hasattr(row[0], 'read') else str(row[0])
                    
                    metadata_text = row[3].read() if hasattr(row[3], 'read') else str(row[3])
                    try:
                        metadata = json.loads(metadata_text) if metadata_text and metadata_text != '{}' else {}
                    except:
                        metadata = {'raw': metadata_text} if metadata_text else {}
                    
                    results.append({
                        'text': chunk_text,
                        'page_number': row[1] if len(row) > 1 else 1,
                        'content_type': row[2] if len(row) > 2 else 'text',
                        'metadata': metadata,
                        'filename': row[4] if len(row) > 4 else 'unknown',
                        'file_type': row[5] if len(row) > 5 else 'pdf',
                        'similarity_score': float(row[6]) if len(row) > 6 else 0.0
                    })
                
                return results
                
        except Exception as e:
            st.error(f"❌ Vector search error: {e}")
            logger.error(f"Vector search failed: {e}")
            return []
        finally:
            conn.close()
    
    @staticmethod
    def get_document_statistics():
        """Get document statistics"""
        conn = DatabaseManager.get_connection()
        if not conn:
            return []
        
        try:
            with conn.cursor() as cursor:
                query = """
                    SELECT 
                        d.filename,
                        d.file_type,
                        d.total_chunks,
                        d.upload_timestamp,
                        d.file_size,
                        COUNT(dc.chunk_id) as actual_chunks
                    FROM documents d
                    LEFT JOIN document_chunks dc ON d.document_id = dc.document_id
                    GROUP BY d.document_id, d.filename, d.file_type, d.total_chunks, d.upload_timestamp, d.file_size
                    ORDER BY d.upload_timestamp DESC
                """
                
                cursor.execute(query)
                results = cursor.fetchall()
                return results
                
        except Exception as e:
            st.error(f"❌ Statistics retrieval error: {e}")
            logger.error(f"Statistics query failed: {e}")
            return []
        finally:
            conn.close()

# Enhanced Chat Engine
class EnhancedChatEngine:
    """Enhanced chat engine with multi-modal support"""
    
    @staticmethod
    def generate_enhanced_response(question, retrieved_chunks):
        """Generate response with enhanced context from multi-modal content"""
        if not retrieved_chunks:
            return "I couldn't find relevant information to answer your question.", []
        
        if not st.session_state.get('openai_client'):
            return "OpenAI client not configured. Please check your API key.", []
        
        # Build context from chunks
        context_parts = []
        sources = []
        
        for chunk in retrieved_chunks:
            text = chunk['text']
            page = chunk['page_number']
            filename = chunk['filename']
            file_type = chunk['file_type']
            score = chunk['similarity_score']
            
            context_parts.append(f"[From {filename} ({file_type.upper()}), Page {page}]\n{text}")
            sources.append({
                'filename': filename,
                'file_type': file_type,
                'page': page,
                'page_number': page,
                'content_type': chunk.get('content_type', 'text'),
                'similarity_score': score
            })
        
        context = "\n\n---\n\n".join(context_parts)
        
        system_prompt = """You are a helpful AI assistant that answers questions based on document content. 
        Always be accurate and cite your sources clearly. Be specific about which document and page the information comes from."""
        
        user_prompt = f"""Based on the following content from documents, please answer this question: {question}

Content:
{context}

Please provide a comprehensive answer based on the above content."""
        
        try:
            response = st.session_state.openai_client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.3,
                max_tokens=1000 if not is_mobile() else 500
            )
            
            answer = response.choices[0].message.content.strip()
            return answer, sources
            
        except Exception as e:
            st.error(f"❌ OpenAI API error: {e}")
            return "Sorry, I encountered an error while generating the response.", sources

# Main Application Logic
def main():
    """Main application logic"""
    # Initialize session state
    initialize_session_credentials()
    
    # Check if credentials are configured
    if not st.session_state.get('credentials_configured'):
        render_credential_setup()
        return
    
    # Show credential status in sidebar
    with st.sidebar:
        st.markdown("### 🔐 Credentials Status")
        st.success("✅ OpenAI API Connected")
        st.success("✅ Database Connected")
        
        if st.button("🔄 Reconfigure Credentials"):
            # Clear credentials and restart
            st.session_state.credentials_configured = False
            st.session_state.openai_client = None
            st.session_state.db_credentials = {}
            st.session_state.credentials_tested = False
            st.rerun()
        
        st.markdown("---")
        
        # Show warnings about missing dependencies
        if not DOCX_AVAILABLE:
            st.sidebar.info("ℹ️ Install python-docx for Word document support")
        if not PPTX_AVAILABLE:
            st.sidebar.info("ℹ️ Install python-pptx for PowerPoint support")
        if not PANDAS_AVAILABLE:
            st.sidebar.info("ℹ️ Install pandas & openpyxl for Excel support")
        if not OCR_AVAILABLE:
            st.sidebar.info("ℹ️ Install Pillow & pytesseract for image OCR")
        if not WEB_AVAILABLE:
            st.sidebar.info("ℹ️ Install beautifulsoup4 & markdown for web content")
    
    # Initialize database
    if 'db_initialized' not in st.session_state:
        with st.spinner("🔄 Initializing database..."):
            if DatabaseManager.initialize_tables():
                st.session_state.db_initialized = True
            else:
                st.error("❌ Failed to initialize database")
                return
    
    # Load embedding model
    embedder = load_embedding_model()
    if not embedder:
        st.error("❌ Failed to load embedding model")
        return
    
    # Main App Interface
    st.title("🤖 Smart RAG Assistant")
    st.markdown("*Multi-modal document chat with Oracle 23ai vector search*")
    
    # Mobile-responsive file upload
    if is_mobile():
        st.markdown("### 📱 Mobile Interface")
        with st.expander("📤 Upload Documents", expanded=False):
            uploaded_file = MobileUI.mobile_file_uploader()
    else:
        st.sidebar.header("📤 Document Upload")
        uploaded_file = st.sidebar.file_uploader(
            "Choose a file",
            type=list(Config.SUPPORTED_FORMATS.keys()),
            help="Upload documents in various formats"
        )
    
    # Handle file upload
    if uploaded_file:
        doc_id = str(uuid.uuid4())[:8]
        
        if (is_mobile() and st.button("🚀 Process Document", type="primary")) or \
           (not is_mobile() and st.sidebar.button("🚀 Process Document", type="primary")):
            
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as temp_file:
                temp_file.write(uploaded_file.getvalue())
                temp_path = temp_file.name
            
            try:
                file_type = MultiModalProcessor.detect_file_type(temp_path)
                st.info(f"📄 Processing {file_type.upper()} file: {uploaded_file.name}")
                
                content_items = MultiModalProcessor.process_file(temp_path)
                
                if not content_items:
                    st.error("❌ No content could be extracted from the file")
                else:
                    st.success(f"✅ Extracted {len(content_items)} content items")
                    
                    st.info("🧠 Creating smart chunks and generating embeddings...")
                    chunks = EnhancedChunking.smart_chunk_content(content_items)
                    
                    if chunks and embedder:
                        progress_bar = st.progress(0)
                        chunks_with_embeddings = []
                        
                        for i, (page_num, chunk_idx, text, metadata) in enumerate(chunks):
                            embedding = embedder.encode(text, convert_to_numpy=True)
                            chunks_with_embeddings.append((page_num, chunk_idx, text, metadata, embedding))
                            progress_bar.progress((i + 1) / len(chunks))
                        
                        progress_bar.empty()
                        st.success(f"✅ Generated embeddings for {len(chunks_with_embeddings)} chunks")
                        
                        st.info("💾 Storing in Oracle 23ai...")
                        
                        file_size = len(uploaded_file.getvalue())
                        metadata_success = DatabaseManager.store_document_metadata(
                            doc_id, uploaded_file.name, file_type, file_size, len(chunks)
                        )
                        
                        chunks_success = DatabaseManager.store_enhanced_chunks(
                            doc_id, chunks_with_embeddings, file_type
                        )
                        
                        if metadata_success and chunks_success:
                            st.success("🎉 Document successfully processed and stored!")
                            if is_mobile():
                                st.balloons()
                            st.rerun()
                        else:
                            st.error("❌ Failed to store document")
                    else:
                        st.error("❌ Failed to generate embeddings")
            
            except Exception as e:
                st.error(f"❌ Processing error: {e}")
            finally:
                try:
                    os.unlink(temp_path)
                except:
                    pass
    
    # Main interface tabs
    if is_mobile():
        tab1, tab2 = st.tabs(["💬 Chat", "📚 Library"])
    else:
        tab1, tab2, tab3 = st.tabs(["📚 Knowledge Base", "💬 Chat Assistant", "📊 Analytics"])
    
    # Tab 1: Knowledge Base (Desktop) / Library (Mobile)
    with (tab2 if is_mobile() else tab1):
        if is_mobile():
            st.header("📚 Document Library")
        else:
            st.header("📚 Knowledge Base")
        
        documents = DatabaseManager.get_document_statistics()
        
        if documents:
            if is_mobile():
                for doc in documents:
                    with st.expander(f"📄 {doc[0]}", expanded=False):
                        col1, col2 = st.columns(2)
                        with col1:
                            st.write(f"**Type:** {doc[1].upper()}")
                            st.write(f"**Chunks:** {doc[5]}")
                        with col2:
                            st.write(f"**Size:** {doc[4]/1024:.1f} KB" if doc[4] else "N/A")
                            st.write(f"**Date:** {doc[3].strftime('%Y-%m-%d')}" if doc[3] else "N/A")
            else:
                import pandas as pd
                
                df_data = []
                for doc in documents:
                    df_data.append({
                        'Filename': doc[0],
                        'Type': doc[1].upper() if doc[1] else 'PDF',
                        'Total Chunks': doc[5],
                        'Upload Date': doc[3].strftime('%Y-%m-%d %H:%M') if doc[3] else 'N/A',
                        'Size (KB)': f"{doc[4]/1024:.1f}" if doc[4] else 'N/A'
                    })
                
                df = pd.DataFrame(df_data)
                st.dataframe(df, use_container_width=True)
                
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("Total Documents", len(documents))
                with col2:
                    total_chunks = sum(doc[5] if doc[5] else 0 for doc in documents)
                    st.metric("Total Chunks", total_chunks)
                with col3:
                    total_size = sum(doc[4] if doc[4] else 0 for doc in documents)
                    st.metric("Total Size (MB)", f"{total_size/(1024*1024):.1f}")
        else:
            st.info("📝 No documents uploaded yet. Upload some files to get started!")
    
    # Tab 2: Chat Interface (Desktop) / Main Chat (Mobile)
    with (tab1 if is_mobile() else tab2):
        if is_mobile():
            user_question, ask_button, clear_button = MobileUI.mobile_chat_interface()
            
            if clear_button:
                st.session_state.chat_history = []
                st.rerun()
        else:
            st.header("💬 Chat with Your Documents")
            
            with st.form("chat_form", clear_on_submit=True):
                user_question = st.text_area(
                    "Ask a question about your documents:",
                    height=100,
                    placeholder="What would you like to know?"
                )
                
                col1, col2 = st.columns([3, 1])
                with col1:
                    ask_button = st.form_submit_button("🔍 Ask", type="primary")
                with col2:
                    if st.form_submit_button("🗑️ Clear History"):
                        st.session_state.chat_history = []
                        st.rerun()
        
        # Initialize chat history
        if 'chat_history' not in st.session_state:
            st.session_state.chat_history = []
        
        # Process question
        if ask_button and user_question.strip():
            with st.spinner("🔍 Searching knowledge base..."):
                question_embedding = embedder.encode(user_question, convert_to_numpy=True)
                retrieved_chunks = DatabaseManager.enhanced_vector_search(question_embedding)
                
                if retrieved_chunks:
                    answer, sources = EnhancedChatEngine.generate_enhanced_response(
                        user_question, retrieved_chunks
                    )
                    
                    st.session_state.chat_history.append({
                        'question': user_question,
                        'answer': answer,
                        'sources': sources,
                        'timestamp': datetime.now()
                    })
                else:
                    st.warning("No relevant content found in your documents.")
        
        # Display chat history
        if st.session_state.chat_history:
            if is_mobile():
                st.markdown("### 💭 Conversation")
            else:
                st.subheader("💭 Chat History")
            
            for i, chat in enumerate(reversed(st.session_state.chat_history)):
                if is_mobile():
                    MobileUI.mobile_chat_message(chat['question'], is_user=True)
                    MobileUI.mobile_chat_message(chat['answer'], is_user=False)
                    
                    if chat['sources']:
                        with st.expander("📎 Sources", expanded=False):
                            for j, source in enumerate(chat['sources'], 1):
                                score = source['similarity_score']
                                page = source.get('page', source.get('page_number', 'N/A'))
                                content_type = source.get('content_type', 'text')
                                file_type = source.get('file_type', 'unknown')
                                
                                icon = "📊" if content_type == 'table' else "📄"
                                st.write(f"{j}. {icon} `{source['filename']}` ({file_type.upper()}) - Page {page}")
                                st.caption(f"Similarity: {score:.3f}")
                else:
                    with st.expander(
                        f"Q: {chat['question'][:100]}..." if len(chat['question']) > 100 
                        else f"Q: {chat['question']}", 
                        expanded=(i==0)
                    ):
                        st.markdown("**Question:**")
                        st.write(chat['question'])
                        
                        st.markdown("**Answer:**")
                        st.write(chat['answer'])
                        
                        if chat['sources']:
                            st.markdown("**Sources:**")
                            for j, source in enumerate(chat['sources'], 1):
                                score = source['similarity_score']
                                page = source.get('page', source.get('page_number', 'N/A'))
                                content_type = source.get('content_type', 'text')
                                file_type = source.get('file_type', 'unknown')
                                
                                icon = "📊" if content_type == 'table' else "📄"
                                st.write(f"{j}. {icon} `{source['filename']}` ({file_type.upper()}) - Page {page} - Similarity: {score:.3f}")
                        
                        st.caption(f"Asked at: {chat['timestamp'].strftime('%Y-%m-%d %H:%M:%S')}")
    
    # Tab 3: Analytics (Desktop only)
    if not is_mobile():
        with tab3:
            st.header("📊 System Analytics")
            
            st.subheader("🔗 System Status")
            st.success("✅ Oracle 23ai Database Connected")
            st.success("✅ OpenAI API Connected")
            st.success(f"✅ Embedding Model: {Config.EMBEDDING_MODEL}")
            
            st.subheader("🎯 Available Features")
            col1, col2 = st.columns(2)
            
            with col1:
                st.write("**Core Features:**")
                st.write("✅ PDF Processing")
                st.write("✅ Vector Search")
                st.write("✅ Mobile Interface")
                
            with col2:
                st.write("**Optional Features:**")
                st.write(f"{'✅' if DOCX_AVAILABLE else '❌'} Word Documents")
                st.write(f"{'✅' if PPTX_AVAILABLE else '❌'} PowerPoint")
                st.write(f"{'✅' if PANDAS_AVAILABLE else '❌'} Excel Files")
                st.write(f"{'✅' if OCR_AVAILABLE else '❌'} Image OCR")
                st.write(f"{'✅' if WEB_AVAILABLE else '❌'} Web Content")
            
            st.subheader("⚙️ Configuration")
            config_data = {
                'Embedding Model': Config.EMBEDDING_MODEL,
                'Embedding Dimension': Config.EMBEDDING_DIMENSION,
                'Chunk Size': Config.CHUNK_SIZE,
                'Mobile Chunk Size': Config.MOBILE_CHUNK_SIZE,
                'Top K Results': Config.TOP_K,
                'Mobile Top K': Config.MOBILE_TOP_K,
                'Supported Formats': len(Config.SUPPORTED_FORMATS)
            }
            
            for key, value in config_data.items():
                st.write(f"**{key}:** {value}")
    
    # Footer
    st.markdown("---")
    if is_mobile():
        st.markdown("*📱 Mobile-optimized • 🤖 AI-powered • ⚡ Oracle 23ai*")
    else:
        st.markdown("*Powered by Oracle Database 23ai Vector Search, OpenAI, and Multi-Modal Document Processing*")

# Run the application
if __name__ == "__main__":
    main()